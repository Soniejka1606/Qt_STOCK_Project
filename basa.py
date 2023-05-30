import sqlite3 as sl
from docxtpl import DocxTemplate
from datetime import date
import pdfkit
from docx import Document
from docx.shared import Inches
import subprocess
from openpyxl import load_workbook
from openpyxl.workbook import Workbook

"""
SELECT ('столбцы или * для выбора всех столбцов; обязательно')
FROM ('таблица; обязательно')
WHERE ('условие/фильтрация, например, city = 'Moscow'; необязательно')
GROUP BY ('столбец, по которому хотим сгруппировать данные; необязательно')
HAVING ('условие/фильтрация на уровне сгруппированных данных; необязательно')
ORDER BY ('столбец, по которому хотим отсортировать вывод; необязательно')
"""

con = sl.connect('DARASTOCK.db', check_same_thread=False)

with con:
    con.execute("""
        CREATE TABLE IF NOT EXISTS Stock (
            id INTEGER NOT NULL PRIMARY KEY AUTOINCREMENT,
            name TEXT,
            address TEXT,
            coordinates REAL);
    """)

with con:
    con.execute("""
        CREATE TABLE IF NOT EXISTS Products (
            id INTEGER NOT NULL PRIMARY KEY AUTOINCREMENT,
            category_id INTEGER,
            name TEXT,
            count INTEGER,
            characteristic TEXT,
            picture BLOB,
            time_out DATETIME,
            stock_id INTEGER,
            articul TEXT);
    """)

with con:
    con.execute("""
        CREATE TABLE IF NOT EXISTS CategoryOfProduct (
            id INTEGER NOT NULL PRIMARY KEY AUTOINCREMENT,
            name TEXT,
            description INTEGER);
    """)

with con:
    con.execute("""
        CREATE TABLE IF NOT EXISTS OrderProduct (
            id INTEGER NOT NULL PRIMARY KEY AUTOINCREMENT,
            product_id INTEGER,
            count INTEGER ,
            order_id INTEGER);
    """)

with con:
    con.execute("""
        CREATE TABLE IF NOT EXISTS Orders (
            id INTEGER NOT NULL PRIMARY KEY AUTOINCREMENT,
            date DATETIME,
            stock_id INTEGER,
            user_id INTEGER);
    """)
with con:
    con.execute("""
        CREATE TABLE IF NOT EXISTS Users (
            id INTEGER NOT NULL PRIMARY KEY AUTOINCREMENT,
            name TEXT,
            number_phone INTEGER,
            passport TEXT);
    """)

with con:
    con.execute("""
        CREATE TABLE IF NOT EXISTS Relocation (
            id INTEGER NOT NULL PRIMARY KEY AUTOINCREMENT,
             from_stock_id INTEGER,
             where_stock_id INTEGER,
             comment TEXT,
             who_is_maim TEXT DEFAULT Admin,
             is_done BOOLEAN DEFAULT 0,
            is_order BOOLEAN DEFAULT 0);
    """)

with con:
    con.execute("""
        CREATE TABLE IF NOT EXISTS RelocationProduct (
            id INTEGER NOT NULL PRIMARY KEY AUTOINCREMENT,
            category_id INTEGER,
            name TEXT,
            count INTEGER,
            characteristic TEXT,
            picture BLOB,
            time_out DATETIME,
            stock_id INTEGER,
            articul TEXT);
    """)

def do_order(dict_): #для оформления заказа
    '''
    {"products":[[id_tovar,count,stock_id]],sklad_end : 1}
    возвращает id заказа
    '''
    relocate = False
    #создвние перемещения
    try:
        sql_insert = f"INSERT INTO Relocation (where_stock_id,is_order) values(?,?)"
        with con:
            con.execute(sql_insert, (str(dict_["sklad_end"]),"1"))
    except Exception as e:
        return print(e)
    # получение id созданного перемещения перемещения
    try:
        order_id = con.execute(f"SELECT id FROM Relocation ORDER BY id DESC LIMIT 1")
        order_id = order_id.fetchall()
        id_reloc = order_id[0][0]
        print(f"id перемещения: {id_reloc}")
    except:
        return False
    #создвние Заказа
    try:
        sql_insert = f"INSERT INTO Orders (user_id,stock_id) values(?,?)"
        with con:
            con.execute(sql_insert, (str(dict_["user_id"]),dict_["sklad_end"]))
    except Exception as e:
        return print(e)
    # получение id созданного Заказа
    try:
        order_id = con.execute(f"SELECT id FROM Orders ORDER BY id DESC LIMIT 1")
        order_id = order_id.fetchall()
        id_order = order_id[0][0]
        print(f'id заказа: {id_order}')
    except:
        return False
    for j in dict_["products"]:
        if j[2]==dict_["sklad_end"]: #если перемещать не надо
            try:
                #списываем товар
                com = f"UPDATE Products SET count = count - ? WHERE Id = ?"
                with con:
                    con.execute(com, (str(j[1]), str(j[0])))
            except Exception as e:
                print("Ошибка: ", e)

            try:
                #добавляем в заказанные товары
                sql_insert = f"INSERT INTO OrderProduct (product_id,count,order_id) values(?,?,?)"
                with con:
                    con.execute(sql_insert, (j[0],j[1],str(id_order)))
            except Exception as e:
                return print(e)
        else: #если надо перемещать
            relocate = True
            try:
                #списываем товар
                com = f"UPDATE Products SET count = count - ? WHERE Id = ?"
                with con:
                    con.execute(com, (str(j[1]), str(j[0])))
            except Exception as e:
                print("Ошибка: ", e)
            try:
                #добавляем в перемещенные товары товары
                sql_insert = f"INSERT INTO RelocationProduct (product_id,count,relocation_id,stock_id) values(?,?,?,?)"
                with con:
                    con.execute(sql_insert, (j[0],j[1],str(id_reloc),j[2]))
            except Exception as e:
                return print(e)
            try:
                #добавляем в заказанные товары
                sql_insert = f"INSERT INTO OrderProduct (product_id,count,order_id) values(?,?,?)"
                with con:
                    con.execute(sql_insert, (j[0],j[1],str(id_order)))
            except Exception as e:
                return print(e)
    if relocate:
        return id_order,id_reloc
    else:
        try:
            sql_delete_query = f"""DELETE from Relocation where id = {id_reloc}"""
            with con:
                con.execute(sql_delete_query)
        except Exception as e:
            print(e)
        return id_order



# print(do_order({"products":[["2","3","2"],["9","2","1"]],"sklad_end" : "1","user_id":"1"}))


def generate_doc_about_order(order_id): # не учитывая перемещения
    '''
    генерируется документ продажи
    :param order_id: айди заказа из прошлой функции
    :return:
    '''
    template = DocxTemplate('docs_about_ordering/шаблон.docx') #пишем где находится шаблон
    try:
        with con:
            data = con.execute(f'''SELECT Users.name FROM Users
                                Join Orders On
                                Users.id = Orders.user_id
                                WHERE Orders.id = "{str(order_id)}"''')
            data = data.fetchall()
            user_name = data[0][0]
        # return user_name
    except Exception as e:
        return e

    today = date.today()
    format_today = today.strftime("%d.%m.%Y")
    products =  []
    try:
        with con:
            data = con.execute(f'''SELECT Products.name,OrderProduct.count,Stock.name,Stock.address FROM OrderProduct
                                Join Products On
                                OrderProduct.product_id = Products.id
                                Join Orders On
                                OrderProduct.order_id = Orders.id
                                Join Stock On
                                Orders.stock_id = Stock.id
                                WHERE Orders.id = "{str(order_id)}"''')
            data = data.fetchall()
            for i in data:
                a = {}
                a["name"] = i[0]
                a["count"] = i[1]
                a["sklad"] = f'{i[2]} {i[3]}'
                products.append(a)
        # return user_name
    except Exception as e:
        return e
    context = {
        'order_id':order_id,
        'name': user_name,
        'date': format_today,
        'orders': products}
    # print(products)
    template.render(context) #вставляем в шаблон нужжные данные
    peczat = 'docs_about_ordering/печать.png' #пишем где находится печать
    do_peczat = template.add_paragraph() #добавляем абзац на добавлеине печати
    do_peczat.add_run().add_picture(peczat, width=Inches(2)) #добавляем печать
    template.save(f'all_acts/Заказа{order_id}.docx') #сохраняем файл
    subprocess.call(['start', f'all_acts/Заказа{order_id}.docx'], shell=True) #для открытия файла на просмотр
    #excel
    template_path = 'docs_about_ordering/шаблон.xlsx'
    workbook = load_workbook(template_path)
    sheet = workbook['Лист1']
    sheet['C2'] = order_id
    sheet['C3'] = format_today
    sheet['C4'] = user_name
    row_index = 7
    for item in products:
        sheet.cell(row=row_index, column=1, value=item['name'])
        sheet.cell(row=row_index, column=2, value=item['count'])
        sheet.cell(row=row_index, column=3, value=item['sklad'])
        row_index += 1
    output_path = f'all_acts/Заказа{order_id}.xlsx'
    workbook.save(output_path)
    subprocess.call(['start', f'all_acts/Заказа{order_id}.xlsx'], shell=True)  # для открытия файла на просмотр

# print(generate_doc_about_order(8))
def generate_doc_about_order_relocat(two_id): # учитывая перемещения
    '''
    генерируется документ что куда откуда надо быстренько привезди для заказа
    :param order_id: айди перемещения из прошлой функции (9, 12)
    :return:
    '''
    template = DocxTemplate('docs_about_ordering/шаблон1.docx')
    try:
        with con:
            data = con.execute(f'''SELECT Stock.name,Stock.address FROM Relocation
            JOIN Stock ON
            Relocation.where_stock_id = Stock.id
            WHERE Relocation.id = {two_id[1]}
                                    ''')
            data = data.fetchall()
            where_address = f'{data[0][0]} {data[0][1]}'

    except Exception as e:
        return e
    today = date.today()
    format_today = today.strftime("%d.%m.%Y")
    products =  []
    try:
        with con:
            data = con.execute(f'''SELECT Products.name,RelocationProduct.count,Stock.name,Stock.address FROM RelocationProduct
            JOIN Products ON
            RelocationProduct.product_id = Products.id
            JOIN Stock ON
            RelocationProduct.stock_id = Stock.id
            JOIN Relocation ON
            RelocationProduct.relocation_id = Relocation.id
            WHERE Relocation.id = {two_id[1]}

                                    ''')
            data = data.fetchall()
            for i in data:
                a = {}
                a["name"] = i[0]
                a["count"] = i[1]
                a["sklad_from"] = f'{i[2]} {i[3]}'
                a["sklad_in"] = where_address
                products.append(a)
            # print(products)


    except Exception as e:
        return e

    context = {
        'order_id':two_id[0],
        'relocate_id':two_id[1],
        'date': format_today,
        'orders': products}

    template.render(context)
    peczat = 'docs_about_ordering/печать.png'
    do_peczat = template.add_paragraph()
    do_peczat.add_run().add_picture(peczat, width=Inches(2))
    template.save(f'all_acts/Перемещение_Заказа{two_id[0]}.docx') #nned
    subprocess.call(['start', f'all_acts/Перемещение_Заказа{two_id[0]}.docx'], shell=True)  # для открытия файла на просмотр
    #exel
    template_path = 'docs_about_ordering/шаблон1.xlsx'
    workbook = load_workbook(template_path)
    sheet = workbook['Лист1'] #какой лист для шаблона
    sheet['C2'] = two_id[0]
    sheet['C3'] = two_id[1]
    sheet['C4'] = format_today
    row_index = 8
    for item in products:
        sheet.cell(row=row_index, column=1, value=item['name'])
        sheet.cell(row=row_index, column=2, value=item['count'])
        sheet.cell(row=row_index, column=3, value=item['sklad_from'])
        sheet.cell(row=row_index, column=4, value=item['sklad_in'])
        row_index += 1
    output_path = f'all_acts/Перемещение_Заказа{two_id[0]}.xlsx'
    workbook.save(output_path)
    subprocess.call(['start', f'all_acts/Перемещение_Заказа{two_id[0]}.xlsx'], shell=True)  # для открытия файла на просмотр
# print(generate_doc_about_order_relocat((10, 13)))


def show_category():
    '''
    :return: список категорий
    '''
    try:
        with con:
            data = con.execute(f"SELECT name FROM CategoryOfProduct")
            data = data.fetchall()
            list_cat = []
            for i in data:
                list_cat.append(i[0])
        return list_cat
    except Exception as e:
        print(e)
# print(show_category())

def show_products():
    '''
    param
    :return: список продуктов со всем нужным
    '''
    try:
        with con:
            data = con.execute(f'''SELECT Products.id,CategoryOfProduct.name,Products.name, 
                                    Products.count,Stock.name,Products.articul,Products.characteristic
                                    FROM Products
                                    JOIN CategoryOfProduct ON
                                    Products.category_id = CategoryOfProduct.id
                                    JOIN Stock ON
                                    Products.stock_id = Stock.id
                                    WHERE Products.count>0
                                     ''')
            data = data.fetchall()
        return data
    except Exception as e:
        print(e)
# print(show_products())


def registr_new_stock(name,address):
    '''
    :param name: название склада и адресс склада
    :param address:
    :return: новый зарегистрированный склад
    '''
    try:
        sql_insert = f"INSERT INTO Stock (name,address) values(?,?)"
        with con:
            con.execute(sql_insert, (name, address))
        return True
    except Exception as e:
        return e
# print(registr_new_stock("СКЛАД1","Минск,Центральная 8"))

def registr_new_user(name,number_phone,passport):
    '''
    :param name: название склада и адресс склада
    :param address:
    :return: новый зарегистрированный склад
    '''
    try:
        sql_insert = f"INSERT INTO Users (name,number_phone,passport) values(?,?,?)"
        with con:
            con.execute(sql_insert, (name, number_phone,passport))
        return True
    except Exception as e:
        return e
# print(registr_new_user("Иван Иванович Иванов",3752922287,"AB656565"))
def show_users():
    '''
    param
    :return: список покупателей
    '''
    try:
        with con:
            data = con.execute(f'''SELECT id,name,number_phone,passport
                                        FROM Users
                                         ''')
            data = data.fetchall()
        return data
    except Exception as e:
        print(e)
# print(show_users())





def add_category(name,descriprion):
    try:
        sql_insert = f"INSERT INTO CategoryOfProduct (name,description) values(?,?)"
        with con:
            con.execute(sql_insert, (name, descriprion))
        return True
    except Exception as e:
        return e
# print(add_category("Обои","Красиво"))

def spisanie(list_of_dict):
    '''
    :param dict_: список со словарями с параметрами id товара, count
    :return:
    '''
    today = date.today()
    format_today = today.strftime("%d.%m.%Y")
    products = []
    for i in list_of_dict:
        try:
            # списываем товар
            com = f"UPDATE Products SET count = count - ? WHERE Id = ?"
            with con:
                con.execute(com, (i["count"], i["id_product"]))
        except Exception as e:
            print("Ошибка: ", e)
        try:
            name_product = con.execute(f'''SELECT Products.name,Stock.name FROM Products
                                              Join Stock ON
                                                Products.stock_id = Stock.id
                                                 WHERE Products.id = "{i['id_product']}"''')
            name_product = name_product.fetchall()

            products.append({"name":name_product[0][0],"stock_name":name_product[0][1],"count":i["count"]})
        except Exception as e:
            return e
    print(products)
    template = DocxTemplate('docs_about_ordering/шаблон_списания.docx')
    context = {
        'date': format_today,
        'orders': products}
    # print(products)
    template.render(context)  # вставляем в шаблон нужжные данные
    peczat = 'docs_about_ordering/печать.png'  # пишем где находится печать
    do_peczat = template.add_paragraph()  # добавляем абзац на добавлеине печати
    do_peczat.add_run().add_picture(peczat, width=Inches(2))  # добавляем печать
    template.save(f'all_acts/Списание{format_today}.docx')  # сохраняем файл
    subprocess.call(['start', f'all_acts/Списание{format_today}.docx'], shell=True)  # для открытия файла на просмотр
    # excel
    template_path = 'docs_about_ordering/шаблон_списания.xlsx'
    workbook = load_workbook(template_path)
    sheet = workbook['Лист1']
    sheet['C2'] = format_today
    row_index = 5
    for item in products:
        sheet.cell(row=row_index, column=1, value=item['name'])
        sheet.cell(row=row_index, column=2, value=item['count'])
        sheet.cell(row=row_index, column=3, value=item['stock_name'])
        row_index += 1
    output_path = f'all_acts/Списание{format_today}.xlsx'
    workbook.save(output_path)
    subprocess.call(['start', f'all_acts/Списание{format_today}.xlsx'], shell=True)  # для открытия файла на просмотр
# print(spisanie([{"id_product": 12, "count": 5},{"id_product": 2, "count": 1}]))

def show_stocks():
    '''
    :return:  список со складами
    '''
    list_of_stock = []
    try:
        stock = con.execute(f'''SELECT name,address FROM Stock
                                          ''')
        stock = stock.fetchall()
        for i in stock:
            list_of_stock.append(f'{i[0]} , {i[1]}')
        return list_of_stock
    except Exception as e:
        return e
# print(show_stocks())

def add_new_product(list_of_dict):
    '''
    :param list_of_dict: список со словарями для добавления
    :return:
    '''
    opis_adding =  []
    for i in list_of_dict:
        try:
            sql_insert = f'''INSERT INTO Products (name,category_id,count,characteristic,stock_id,articul,time_out) values(?,?,?,?,?,?,?)'''
            cat_id = con.execute(f'''SELECT id FROM CategoryOfProduct WHERE name = "{i["category"]}"''')
            cat_id = cat_id.fetchall()
            cat_id = cat_id[0][0]
            stock_id = con.execute(f'''SELECT id FROM Stock WHERE name = "{i["sklad"]}"''')
            stock_id = stock_id.fetchall()
            stock_id = stock_id[0][0]
            opis_adding.append({"name":i["name"],"count":i["count"],"sklad":i["sklad"],"category":i["category"]})
            with con:
                con.execute(sql_insert, (i["name"],cat_id,i["count"],i["characteristic"],stock_id,i["articul"],i["time_out"]))
        except Exception as e:
            return print(e)
    #word
    today = date.today()
    format_today = today.strftime("%d.%m.%Y")
    template = DocxTemplate('docs_about_ordering/шаблон_добавления.docx')
    context = {
        'date': format_today,
        'orders': opis_adding}
    # print(products)
    template.render(context)  # вставляем в шаблон нужжные данные
    peczat = 'docs_about_ordering/печать.png'  # пишем где находится печать
    do_peczat = template.add_paragraph()  # добавляем абзац на добавлеине печати
    do_peczat.add_run().add_picture(peczat, width=Inches(2))  # добавляем печать
    template.save(f'all_acts/Добавлеине{format_today}.docx')  # сохраняем файл
    subprocess.call(['start', f'all_acts/Добавлеине{format_today}.docx'], shell=True)  # для открытия файла на просмотр
    # excel
    template_path = 'docs_about_ordering/шаблон_добавления.xlsx'
    workbook = load_workbook(template_path)
    sheet = workbook['Лист1']
    sheet['C2'] = format_today
    row_index = 5
    for item in opis_adding:
        sheet.cell(row=row_index, column=1, value=item['name'])
        sheet.cell(row=row_index, column=2, value=item['count'])
        sheet.cell(row=row_index, column=3, value=item['sklad'])
        sheet.cell(row=row_index, column=4, value=item['category'])
        row_index += 1
    output_path = f'all_acts/Добавление{format_today}.xlsx'
    workbook.save(output_path)
    subprocess.call(['start', f'all_acts/Добавление{format_today}.xlsx'], shell=True)  # для открытия файла на просмотр

# print(add_new_product([{"name":"Кровать Будапешт","category":"Кровати","count":3,"characteristic":'что-то',"sklad":"СКЛАД2","articul":"BED3BUDAPESZT","time_out":"31.08.2025"}]))

def relocate(list_of_dict,end_stock,comment,who_is_maim = 0):
    '''
    только на оформление перемещения пока без принятия
    :param list_of_dict: список со словарями о продуктах
    :return:
    '''
    #создание перемещения
    try:
        stock_id = con.execute(f'''SELECT id FROM Stock WHERE name = "{end_stock}"''')
        stock_id = stock_id.fetchall()
        stock_id = stock_id[0][0]
        if who_is_maim !=0:
            sql_insert = f"INSERT INTO Relocation (where_stock_id,comment,who_is_maim) values(?,?,?)"
            with con:
                con.execute(sql_insert, (stock_id,comment,who_is_maim))
        else:
            sql_insert = f"INSERT INTO Relocation (where_stock_id,comment) values(?,?)"
            with con:
                con.execute(sql_insert, (stock_id, comment))
    except Exception as e:
        return print(e)
    # получение id созданного перемещения перемещения
    try:
        order_id = con.execute(f"SELECT id FROM Relocation ORDER BY id DESC LIMIT 1")
        order_id = order_id.fetchall()
        id_reloc = order_id[0][0]
        print(f"id перемещения: {id_reloc}")
    except:
        return False
    opis_adding=[]
    #добавление перемещнных товаров и списание
    for i in list_of_dict:
        try:
            # списываем товар
            com = f"UPDATE Products SET count = count - ? WHERE Id = ?"
            with con:
                con.execute(com, (i["count"],i["id"]))
        except Exception as e:
            print("Ошибка: ", e)
        try:
            sql_insert = f'''INSERT INTO RelocationProduct (name,category_id,count,characteristic,stock_id,articul,time_out,relocation_id,product_id) values(?,?,?,?,?,?,?,?,?)'''
            cat_id = con.execute(f'''SELECT id FROM CategoryOfProduct WHERE name = "{i["category"]}"''')
            cat_id = cat_id.fetchall()
            cat_id = cat_id[0][0]
            stock_id = con.execute(f'''SELECT id FROM Stock WHERE name = "{i["sklad"]}"''')
            stock_id = stock_id.fetchall()
            stock_id = stock_id[0][0]
            opis_adding.append({"name": i["name"], "count": i["count"], "sklad": i["sklad"], "sklad_in": end_stock})
            with con:
                con.execute(sql_insert,
                            (i["name"], cat_id, i["count"], i["characteristic"], stock_id, i["articul"], i["time_out"],id_reloc,i["id"]))

        except Exception as e:
            return print(e)
        # word
        today = date.today()
        format_today = today.strftime("%d.%m.%Y")
        template = DocxTemplate('docs_about_ordering/шаблон__перемещния.docx')
        context = {
            "id": id_reloc,
            "sklad_in": end_stock,
            'date': format_today,
            'orders': opis_adding}
        # print(products)
        template.render(context)  # вставляем в шаблон нужжные данные
        peczat = 'docs_about_ordering/печать.png'  # пишем где находится печать
        do_peczat = template.add_paragraph()  # добавляем абзац на добавлеине печати
        do_peczat.add_run().add_picture(peczat, width=Inches(2))  # добавляем печать
        template.save(f'all_acts/Перемещение{format_today}.docx')  # сохраняем файл
        subprocess.call(['start', f'all_acts/Перемещение{format_today}.docx'], shell=True)  # для открытия файла на просмотр
        # excel
        template_path = 'docs_about_ordering/шаблон_перемещения.xlsx'
        workbook = load_workbook(template_path)
        sheet = workbook['Лист1']
        sheet['C1'] = id_reloc
        sheet['C2'] = format_today
        sheet['C3'] = end_stock
        row_index = 6
        for item in opis_adding:
            sheet.cell(row=row_index, column=1, value=item['name'])
            sheet.cell(row=row_index, column=2, value=item['count'])
            sheet.cell(row=row_index, column=3, value=item['sklad'])
            row_index += 1
        output_path = f'all_acts/перемещение{format_today}.xlsx'
        workbook.save(output_path)
        subprocess.call(['start', f'all_acts/перемещение{format_today}.xlsx'], shell=True)  # для открытия файла на просмотр

#print(relocate([{"id":17,"name":"Кровать Будапешт","category":"Кровати","count":3,"characteristic":'что-то',"sklad":"СКЛАД2","articul":"BED3BUDAPESZT","time_out":"31.08.2025"}],"СКЛАД3","переметить"))

def accept_relocate(id_relocate):
    '''
    подтверждает перемещение
    :param id_relocate: айдщи перемещения
    :return:
    '''
    products = []
    try:
        order_id = con.execute(f'''SELECT * FROM RelocationProduct
                                JOIN Relocation ON
                                RelocationProduct.relocation_id = Relocation.id
                                    WHERE Relocation.id = {id_relocate} ''')
        order_id = order_id.fetchall()
        for i in order_id:
            products.append({"name":i[1],"count":i[2],"characteristic":i[3],"time_out":i[5],"stock_from":i[6],"articul":i[7],"category_id":i[8],"stock_in":i[13]})
        print(products)
    except:
        return False
    for i in products:
        try:
            sql_insert = f'''INSERT INTO Products (name,category_id,count,characteristic,stock_id,articul,time_out) values(?,?,?,?,?,?,?)'''
            with con:
                con.execute(sql_insert,
                            (i["name"], i["category_id"], i["count"], i["characteristic"], i["stock_in"], i["articul"], i["time_out"]))
        except Exception as e:
            return print(e)
# print(accept_relocate(19))

def read_doc():
    '''
    попытки читать с документов
    :return:
    '''
    doc = Document('all_acts/Перемещение29.05.2023.docx')
    first_paragraph = doc.paragraphs[0]
    text = first_paragraph.text
    print(text)
    template_path = 'all_acts/перемещение29.05.2023.xlsx'
    workbook = load_workbook(template_path)
    sheet = workbook['Лист1']
    text = sheet["C1"].value
    print(text)
    for row in sheet.iter_rows(min_row=6, min_col=0):
        for cell in row:
            value = cell.value
            print(value)
# print(read_doc())